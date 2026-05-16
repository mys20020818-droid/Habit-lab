from __future__ import annotations

import math
import shutil
import warnings
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import scipy.stats as st
import seaborn as sns
import statsmodels.api as sm
import statsmodels.formula.api as smf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from statsmodels.stats.multitest import multipletests

from s32_fig5_erp_waveform_topomap_grid import highlight_roi, plot_topomap


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "results/supplementary_materials_final"
TABLES = OUT / "tables"
FIGURES = OUT / "figures"
TEXT = OUT / "text"
RNG = np.random.default_rng(20260512)

AGE_GROUPS = ["child_5_9", "adolescent_10_14", "youth_15_21"]
AGE_LABELS = {"child_5_9": "5-9", "adolescent_10_14": "10-14", "youth_15_21": "15-21"}
FEATURES = ["p1", "n1", "p3"]
FEATURE_LABELS = {"p1": "P1", "n1": "P2", "p3": "N450"}
COMPONENTS = ["component_general", "component_specific_sus_minus_ccd"]
COMPONENT_LABELS = {"component_general": "General", "component_specific_sus_minus_ccd": "Specific"}
PSYCH = ["p_factor", "attention", "internalizing", "externalizing"]
PSYCH_LABELS = {
    "p_factor": "p-factor",
    "attention": "Attention",
    "internalizing": "Internalizing",
    "externalizing": "Externalizing",
}
AGE_COLORS = {"child_5_9": "#6DA6D9", "adolescent_10_14": "#D9A441", "youth_15_21": "#B64342"}
COMP_COLORS = {"component_general": "#3F3F3F", "component_specific_sus_minus_ccd": "#7462D8"}


def configure_style() -> None:
    sns.set_theme(style="ticks", context="paper")
    plt.rcParams.update(
        {
            "font.family": "sans-serif",
            "font.sans-serif": ["Arial", "DejaVu Sans", "Liberation Sans"],
            "font.size": 7,
            "axes.titlesize": 8,
            "axes.labelsize": 7,
            "xtick.labelsize": 6.5,
            "ytick.labelsize": 6.5,
            "legend.fontsize": 6.5,
            "axes.linewidth": 0.75,
            "axes.spines.right": False,
            "axes.spines.top": False,
            "svg.fonttype": "none",
            "pdf.fonttype": 42,
        }
    )


def ensure_dirs() -> None:
    for path in [OUT, TABLES, FIGURES, TEXT]:
        path.mkdir(parents=True, exist_ok=True)


def save_figure(fig: plt.Figure, base: Path) -> None:
    base.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(base.with_suffix(".svg"), bbox_inches="tight")
    fig.savefig(base.with_suffix(".pdf"), bbox_inches="tight")
    fig.savefig(base.with_suffix(".png"), dpi=300, bbox_inches="tight")
    plt.close(fig)


def fmt_p(value: float) -> str:
    if pd.isna(value):
        return "NA"
    return "<0.001" if value < 0.001 else f"{value:.3f}"


def zscore_by_task(df: pd.DataFrame, value_col: str = "amplitude_uv") -> pd.DataFrame:
    out = df.copy()
    out[f"{value_col}_z"] = out.groupby("task")[value_col].transform(lambda s: (s - s.mean()) / s.std(ddof=1))
    return out


def age_group_order(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["age_group"] = pd.Categorical(out["age_group"], categories=AGE_GROUPS, ordered=True)
    return out


def load_components() -> pd.DataFrame:
    components = pd.read_csv(ROOT / "results/tables/full_sus_ccd_components.csv")
    for col in ["age", *PSYCH]:
        components[col] = pd.to_numeric(components[col], errors="coerce")
    return age_group_order(components)


def load_features() -> pd.DataFrame:
    features = pd.read_csv(ROOT / "results/tables/full_sus_ccd_erp_features_by_subject_task.csv")
    features = features[features["task"].isin(["CCD", "SuS"])].copy()
    features["amplitude_uv"] = pd.to_numeric(features["amplitude_uv"], errors="coerce")
    return features


def write_table_s1() -> Path:
    sources = {
        "Full sample": ROOT / "results/tables/full_sus_ccd_age_group_anova.csv",
        "High-quality EEG": ROOT / "results/tables/full_sus_ccd_followup/full_sus_ccd_high_quality_age_group_anova.csv",
        "Complete phenotype": ROOT / "results/tables/full_sus_ccd_followup/full_sus_ccd_full_pheno_age_group_anova.csv",
        "High-quality + phenotype": ROOT / "results/tables/full_sus_ccd_followup/full_sus_ccd_high_quality_full_pheno_age_group_anova.csv",
    }
    rows = []
    for label, path in sources.items():
        tab = pd.read_csv(path)
        tab = tab[tab["term"].eq("C(age_group)")].copy()
        tab["sample"] = label
        rows.append(tab)
    out = pd.concat(rows, ignore_index=True)
    out = out[["sample", "feature", "component", "n", "df", "f_value", "p_value", "q_value", "r_squared"]]
    out_path = TABLES / "Supplementary_Table_S1_robustness_all_core_effects.csv"
    out.to_csv(out_path, index=False)
    return out_path


def bootstrap_icc(feature_df: pd.DataFrame, n_boot: int = 1000) -> tuple[float, float, float, float]:
    sub = zscore_by_task(feature_df.dropna(subset=["amplitude_uv"]), "amplitude_uv")
    wide = sub.pivot(index="subject", columns="task", values="amplitude_uv_z").dropna(subset=["CCD", "SuS"])
    if wide.empty:
        return (np.nan, np.nan, np.nan, np.nan)
    subjects = wide.index.to_numpy()
    gen = []
    spec = []
    for _ in range(n_boot):
        sample = RNG.choice(subjects, size=len(subjects), replace=True)
        boot = wide.loc[sample, ["CCD", "SuS"]].reset_index(drop=True)
        stacked = boot.stack()
        subject_mean = boot.mean(axis=1)
        general_repeated = pd.Series(np.repeat(subject_mean.values, 2), index=stacked.index)
        residual = stacked - general_repeated
        total_var = float(stacked.var(ddof=1))
        if total_var <= 0:
            continue
        gen.append(float(general_repeated.var(ddof=1) / total_var))
        spec.append(float(residual.var(ddof=1) / total_var))
    return (
        float(np.nanpercentile(gen, 2.5)),
        float(np.nanpercentile(gen, 97.5)),
        float(np.nanpercentile(spec, 2.5)),
        float(np.nanpercentile(spec, 97.5)),
    )


def fit_mixed_models_for_feature(feature: str, feature_df: pd.DataFrame) -> tuple[list[dict], pd.DataFrame]:
    df = zscore_by_task(feature_df[feature_df["feature"].eq(feature)].copy(), "amplitude_uv")
    df = df.dropna(subset=["amplitude_uv_z", "subject", "task"]).copy()
    df["task_code"] = (df["task"] == "SuS").astype(float)
    rows: list[dict] = []
    diag = pd.DataFrame()

    model_results = {}
    optimizers = ["lbfgs", "bfgs", "powell", "nm", "cg"]
    for re_formula, model_name in [("1", "random_intercept"), ("1 + task_code", "random_intercept_slope")]:
        for reml in [False, True]:
            key = (model_name, "REML" if reml else "ML")
            row = {
                "feature": feature,
                "model": model_name,
                "estimation": "REML" if reml else "ML",
                "n_observations": len(df),
                "n_subjects": df["subject"].nunique(),
                "converged": False,
                "warning_or_error": "",
            }
            try:
                fit = None
                all_warnings: list[str] = []
                last_error: Exception | None = None
                for optimizer in optimizers:
                    try:
                        with warnings.catch_warnings(record=True) as caught:
                            warnings.simplefilter("always")
                            model = smf.mixedlm(
                                "amplitude_uv_z ~ task_code",
                                df,
                                groups=df["subject"],
                                re_formula=re_formula,
                            )
                            candidate = model.fit(reml=reml, method=optimizer, maxiter=500, disp=False)
                        all_warnings.extend(str(w.message) for w in caught)
                        if getattr(candidate, "converged", False):
                            fit = candidate
                            row["optimizer"] = optimizer
                            break
                        fit = candidate
                        row["optimizer"] = optimizer
                    except Exception as opt_exc:
                        last_error = opt_exc
                        all_warnings.append(f"{optimizer}: {type(opt_exc).__name__}: {opt_exc}")
                if fit is None:
                    raise last_error if last_error is not None else RuntimeError("MixedLM fit failed")
                model_results[key] = fit
                warn_text = "; ".join(sorted(set(all_warnings)))
                cov_re = fit.cov_re
                row.update(
                    {
                        "converged": bool(getattr(fit, "converged", False)),
                        "warning_or_error": warn_text,
                        "log_likelihood": float(fit.llf),
                        "aic": float(fit.aic) if not reml else np.nan,
                        "bic": float(fit.bic) if not reml else np.nan,
                        "fixed_intercept_beta": float(fit.params.get("Intercept", np.nan)),
                        "fixed_intercept_se": float(fit.bse.get("Intercept", np.nan)),
                        "fixed_task_sus_vs_ccd_beta": float(fit.params.get("task_code", np.nan)),
                        "fixed_task_sus_vs_ccd_se": float(fit.bse.get("task_code", np.nan)),
                        "subject_intercept_variance": float(cov_re.iloc[0, 0]) if cov_re.shape[0] >= 1 else np.nan,
                        "subject_task_slope_variance": float(cov_re.loc["task_code", "task_code"])
                        if "task_code" in cov_re.index
                        else 0.0,
                        "intercept_slope_covariance": float(cov_re.iloc[0, 1]) if cov_re.shape[0] > 1 else np.nan,
                        "residual_variance": float(fit.scale),
                    }
                )
                total = row["subject_intercept_variance"] + row["subject_task_slope_variance"] + row["residual_variance"]
                row["icc_intercept_lme"] = row["subject_intercept_variance"] / total if total > 0 else np.nan
                row["icc_slope_lme"] = row["subject_task_slope_variance"] / total if total > 0 else np.nan
                if model_name == "random_intercept_slope" and not reml:
                    pred = fit.fittedvalues
                    resid = fit.resid
                    re_rows = []
                    for subject, values in fit.random_effects.items():
                        re_rows.append(
                            {
                                "subject": subject,
                                "feature": feature,
                                "random_intercept": float(values.get("Group", values.iloc[0])),
                                "random_slope": float(values.get("task_code", np.nan)),
                            }
                        )
                    diag = df[["subject", "task", "amplitude_uv_z", "task_code"]].copy()
                    diag["feature"] = feature
                    diag["fitted"] = np.asarray(pred)
                    diag["residual"] = np.asarray(resid)
                    diag = diag.merge(pd.DataFrame(re_rows), on=["subject", "feature"], how="left")
            except Exception as exc:
                row["warning_or_error"] = f"{type(exc).__name__}: {exc}"
            rows.append(row)

    if ("random_intercept", "ML") in model_results and ("random_intercept_slope", "ML") in model_results:
        reduced = model_results[("random_intercept", "ML")]
        full = model_results[("random_intercept_slope", "ML")]
        lr = max(0.0, 2 * (full.llf - reduced.llf))
        # Random slope adds variance and covariance parameters.
        p = st.chi2.sf(lr, df=2)
        for row in rows:
            if row["feature"] == feature:
                row["lr_test_random_slope_vs_intercept"] = lr
                row["lr_test_df"] = 2
                row["lr_test_p_value"] = p

    return rows, diag


def write_table_s2_and_diagnostics() -> tuple[Path, pd.DataFrame]:
    features = load_features()
    rows: list[dict] = []
    diags = []
    ci_rows = {}
    for feature in FEATURES:
        ci_rows[feature] = bootstrap_icc(features[features["feature"].eq(feature)])
        fit_rows, diag = fit_mixed_models_for_feature(feature, features)
        rows.extend(fit_rows)
        if not diag.empty:
            diags.append(diag)
    out = pd.DataFrame(rows)
    for idx, row in out.iterrows():
        g_lo, g_hi, s_lo, s_hi = ci_rows[row["feature"]]
        out.loc[idx, "icc_general_bootstrap_ci_low"] = g_lo
        out.loc[idx, "icc_general_bootstrap_ci_high"] = g_hi
        out.loc[idx, "icc_specific_bootstrap_ci_low"] = s_lo
        out.loc[idx, "icc_specific_bootstrap_ci_high"] = s_hi
    out_path = TABLES / "Supplementary_Table_S2_LME_variance_decomposition_audit.csv"
    out.to_csv(out_path, index=False)
    diag_df = pd.concat(diags, ignore_index=True) if diags else pd.DataFrame()
    diag_df.to_csv(TABLES / "Supplementary_Table_S2_model_diagnostics_source.csv", index=False)
    return out_path, diag_df


def write_table_s3() -> Path:
    src = ROOT / "results/tables/full_sus_ccd_followup/full_sus_ccd_age_group_posthoc.csv"
    out = TABLES / "Supplementary_Table_S3_all_posthoc_pairwise_comparisons.csv"
    shutil.copy2(src, out)
    return out


def spearman_brown(r: float) -> float:
    if pd.isna(r) or r <= -1:
        return np.nan
    return (2 * r) / (1 + r)


def bootstrap_reliability(df: pd.DataFrame, x: str, y: str, n_boot: int = 1000) -> tuple[float, float]:
    valid = df.dropna(subset=[x, y]).copy()
    if len(valid) < 5:
        return np.nan, np.nan
    vals = []
    idx = np.arange(len(valid))
    for _ in range(n_boot):
        sample = valid.iloc[RNG.choice(idx, size=len(idx), replace=True)]
        r = sample[x].corr(sample[y])
        vals.append(spearman_brown(r))
    return float(np.nanmean(vals)), float(np.nanstd(vals, ddof=1))


def write_table_s4() -> Path:
    features = pd.read_csv(ROOT / "results/tables/full_sus_ccd_erp_features_by_subject_task.csv")
    components = load_components()
    rows = []
    for (task, feature), df in features.groupby(["task", "feature"]):
        valid = df.dropna(subset=["split_odd_uv", "split_even_uv"])
        r = valid["split_odd_uv"].corr(valid["split_even_uv"])
        mean_sb, sd_sb = bootstrap_reliability(valid, "split_odd_uv", "split_even_uv")
        rows.append(
            {
                "estimate_family": "task_level_raw_erp",
                "task": task,
                "feature": feature,
                "component": "",
                "n": len(valid),
                "raw_split_half_r": r,
                "spearman_brown": spearman_brown(r),
                "resampling_mean_spearman_brown": mean_sb,
                "resampling_sd_spearman_brown": sd_sb,
                "resampling_note": "1000 subject-level bootstrap resamples of the available odd/even split",
            }
        )
    for feature, df in components.groupby("feature"):
        for component, odd, even in [
            ("general", "general_odd", "general_even"),
            ("specific_sus_minus_ccd", "specific_odd", "specific_even"),
        ]:
            valid = df.dropna(subset=[odd, even])
            r = valid[odd].corr(valid[even])
            mean_sb, sd_sb = bootstrap_reliability(valid, odd, even)
            rows.append(
                {
                    "estimate_family": "decomposed_component",
                    "task": "",
                    "feature": feature,
                    "component": component,
                    "n": len(valid),
                    "raw_split_half_r": r,
                    "spearman_brown": spearman_brown(r),
                    "resampling_mean_spearman_brown": mean_sb,
                    "resampling_sd_spearman_brown": sd_sb,
                    "resampling_note": "1000 subject-level bootstrap resamples of the available odd/even split",
                }
            )
    out = pd.DataFrame(rows)
    out_path = TABLES / "Supplementary_Table_S4_split_half_reliability_complete.csv"
    out.to_csv(out_path, index=False)
    return out_path


def partial_r_from_t(t_value: float, df_resid: float) -> float:
    if pd.isna(t_value) or pd.isna(df_resid):
        return np.nan
    return float(np.sign(t_value) * math.sqrt((t_value * t_value) / (t_value * t_value + df_resid)))


def write_table_s5() -> tuple[Path, Path]:
    components = load_components()
    assoc_rows = []
    interaction_rows = []
    for feature in FEATURES:
        fdf = components[components["feature"].eq(feature)].copy()
        fdf["age_c"] = fdf["age"] - fdf["age"].mean()
        fdf["age2_c"] = fdf["age_c"] ** 2
        for component in COMPONENTS:
            for psych in PSYCH:
                df = fdf.dropna(subset=[component, "age_c", "age2_c", "sex", psych]).copy()
                if len(df) < 30:
                    continue
                df[f"{psych}_c"] = df[psych] - df[psych].mean()
                pred = f"{psych}_c"
                base = smf.ols(f"{component} ~ age_c + age2_c + C(sex) + {pred}", data=df).fit()
                t_value = float(base.tvalues.get(pred, np.nan))
                assoc_rows.append(
                    {
                        "feature": feature,
                        "component": component,
                        "dimension": psych,
                        "n": int(base.nobs),
                        "estimate": float(base.params.get(pred, np.nan)),
                        "std_error": float(base.bse.get(pred, np.nan)),
                        "t_value": t_value,
                        "partial_r": partial_r_from_t(t_value, base.df_resid),
                        "p_value": float(base.pvalues.get(pred, np.nan)),
                        "r_squared": float(base.rsquared),
                    }
                )
                lin = smf.ols(f"{component} ~ age_c + age2_c + C(sex) + {pred} + age_c:{pred}", data=df).fit()
                quad = smf.ols(
                    f"{component} ~ age_c + age2_c + C(sex) + {pred} + age_c:{pred} + age2_c:{pred}",
                    data=df,
                ).fit()
                f_lin, p_lin, df_lin = lin.compare_f_test(base)
                f_quad, p_quad, df_quad = quad.compare_f_test(base)
                interaction_rows.extend(
                    [
                        {
                            "feature": feature,
                            "component": component,
                            "dimension": psych,
                            "test": "age_x_dimension",
                            "n": int(lin.nobs),
                            "df_diff": float(df_lin),
                            "f_value": float(f_lin),
                            "p_value": float(p_lin),
                            "delta_r2": float(lin.rsquared - base.rsquared),
                            "age_x_dimension_beta": float(lin.params.get(f"age_c:{pred}", np.nan)),
                            "age2_x_dimension_beta": np.nan,
                            "model_r_squared": float(lin.rsquared),
                        },
                        {
                            "feature": feature,
                            "component": component,
                            "dimension": psych,
                            "test": "age_and_age2_x_dimension",
                            "n": int(quad.nobs),
                            "df_diff": float(df_quad),
                            "f_value": float(f_quad),
                            "p_value": float(p_quad),
                            "delta_r2": float(quad.rsquared - base.rsquared),
                            "age_x_dimension_beta": float(quad.params.get(f"age_c:{pred}", np.nan)),
                            "age2_x_dimension_beta": float(quad.params.get(f"age2_c:{pred}", np.nan)),
                            "model_r_squared": float(quad.rsquared),
                        },
                    ]
                )
    assoc = pd.DataFrame(assoc_rows)
    inter = pd.DataFrame(interaction_rows)
    assoc["q_value"] = multipletests(assoc["p_value"].fillna(1.0), method="fdr_bh")[1]
    inter["q_value"] = multipletests(inter["p_value"].fillna(1.0), method="fdr_bh")[1]
    assoc_path = TABLES / "Supplementary_Table_S5A_psychopathology_partial_associations_complete.csv"
    inter_path = TABLES / "Supplementary_Table_S5B_psychopathology_age_interactions_complete.csv"
    assoc.to_csv(assoc_path, index=False)
    inter.to_csv(inter_path, index=False)
    return assoc_path, inter_path


def write_table_s6() -> Path:
    iiv = pd.read_csv(ROOT / "results/tables/clinical_rescue_iiv/iiv_components.csv")
    for col in ["age", *PSYCH]:
        iiv[col] = pd.to_numeric(iiv[col], errors="coerce")
    rows = []
    for feature in FEATURES:
        fdf = iiv[iiv["feature"].eq(feature)].copy()
        fdf["age_c"] = fdf["age"] - fdf["age"].mean()
        fdf["age2_c"] = fdf["age_c"] ** 2
        for outcome in ["iiv_general", "iiv_specific_sus_minus_ccd"]:
            for psych in PSYCH:
                df = fdf.dropna(subset=[outcome, "age_c", "age2_c", "sex", psych]).copy()
                if len(df) < 30:
                    continue
                df[f"{psych}_c"] = df[psych] - df[psych].mean()
                pred = f"{psych}_c"
                model = smf.ols(f"{outcome} ~ age_c + age2_c + C(sex) + {pred}", data=df).fit()
                t_value = float(model.tvalues.get(pred, np.nan))
                rows.append(
                    {
                        "feature": feature,
                        "iiv_outcome": outcome,
                        "dimension": psych,
                        "n": int(model.nobs),
                        "estimate": float(model.params.get(pred, np.nan)),
                        "std_error": float(model.bse.get(pred, np.nan)),
                        "t_value": t_value,
                        "partial_r": partial_r_from_t(t_value, model.df_resid),
                        "p_value": float(model.pvalues.get(pred, np.nan)),
                        "r_squared": float(model.rsquared),
                    }
                )
    out = pd.DataFrame(rows)
    out["q_value"] = multipletests(out["p_value"].fillna(1.0), method="fdr_bh")[1]
    out_path = TABLES / "Supplementary_Table_S6_IIV_four_dimension_associations_complete.csv"
    out.to_csv(out_path, index=False)
    return out_path


def write_table_s7() -> tuple[Path, Path]:
    src_a = ROOT / "results/tables/clinical_rescue_behavior/ccd_behavior_pfactor_models.csv"
    src_b = ROOT / "results/tables/clinical_rescue_behavior/ccd_behavior_mediation_models.csv"
    dst_a = TABLES / "Supplementary_Table_S7A_CCD_behavior_pfactor_models.csv"
    dst_b = TABLES / "Supplementary_Table_S7B_CCD_behavior_mediation_models.csv"
    shutil.copy2(src_a, dst_a)
    shutil.copy2(src_b, dst_b)
    return dst_a, dst_b


def write_table_s8() -> Path:
    components = load_components()
    subjects = components[components["feature"].eq("p1")].drop_duplicates("subject").copy()
    qc = pd.read_csv(ROOT / "results/tables/full_sus_ccd_epoch_qc_analysisch_reject250.csv")
    qc = qc[qc["status"].eq("ok")].copy()
    qc["n_epochs_kept"] = pd.to_numeric(qc["n_epochs_kept"], errors="coerce")
    task_epochs = qc.groupby(["subject", "task"])["n_epochs_kept"].sum().unstack("task").reset_index()
    task_epochs = task_epochs.rename(columns={"CCD": "ccd_epochs_kept", "SuS": "sus_epochs_kept"})
    subjects = subjects.merge(task_epochs, on="subject", how="left")
    rows = []
    for group in AGE_GROUPS + ["outside_5_21"]:
        sub = subjects[subjects["age_group"].eq(group)].copy()
        if sub.empty:
            continue
        row = {
            "age_group": AGE_LABELS.get(group, "outside predefined bins"),
            "n_subjects": sub["subject"].nunique(),
            "n_female": int(sub["sex"].eq("F").sum()),
            "n_male": int(sub["sex"].eq("M").sum()),
            "race_ethnicity": "not available in local HBN EEG participants files",
            "n_full_phenotype": int(sub["full_pheno"].eq("Yes").sum()),
            "age_mean": sub["age"].mean(),
            "age_sd": sub["age"].std(),
            "age_min": sub["age"].min(),
            "age_max": sub["age"].max(),
        }
        for var in [*PSYCH, "sus_epochs_kept", "ccd_epochs_kept"]:
            vals = pd.to_numeric(sub[var], errors="coerce")
            row[f"{var}_n"] = vals.notna().sum()
            row[f"{var}_mean"] = vals.mean()
            row[f"{var}_sd"] = vals.std()
            row[f"{var}_min"] = vals.min()
            row[f"{var}_max"] = vals.max()
        rows.append(row)
    out = pd.DataFrame(rows)
    out_path = TABLES / "Supplementary_Table_S8_demographics_and_epoch_counts_complete.csv"
    out.to_csv(out_path, index=False)
    return out_path


def plot_s1_qc() -> Path:
    qc = pd.read_csv(ROOT / "results/tables/full_sus_ccd_epoch_qc_analysisch_reject250.csv")
    meta = load_components()[["subject", "age_group"]].drop_duplicates()
    qc["n_epochs_kept"] = pd.to_numeric(qc["n_epochs_kept"], errors="coerce")
    qc["n_visual_events"] = pd.to_numeric(qc["n_visual_events"], errors="coerce")
    ok = qc[qc["status"].eq("ok")].copy()
    subj_task = ok.groupby(["subject", "task"], as_index=False).agg(
        n_epochs_kept=("n_epochs_kept", "sum"), n_visual_events=("n_visual_events", "sum")
    )
    subj_task["drop_fraction"] = 1 - subj_task["n_epochs_kept"] / subj_task["n_visual_events"]
    subj_task = subj_task.merge(meta, on="subject", how="left")
    subj_task = subj_task[subj_task["age_group"].isin(AGE_GROUPS)].copy()
    fig, axes = plt.subplots(1, 3, figsize=(7.2, 2.25))
    for task, color in [("SuS", "#4E79A7"), ("CCD", "#E15759")]:
        sub = subj_task[subj_task["task"].eq(task)]
        axes[0].hist(sub["n_epochs_kept"], bins=28, alpha=0.55, label=task, color=color)
        axes[1].hist(sub["drop_fraction"] * 100, bins=28, alpha=0.55, label=task, color=color)
    axes[0].set_title("A Retained epochs")
    axes[0].set_xlabel("Epochs per subject-task")
    axes[0].set_ylabel("Subjects")
    axes[1].set_title("B Epoch rejection")
    axes[1].set_xlabel("Discarded epochs (%)")
    sns.boxplot(
        data=subj_task,
        x="age_group",
        y="n_epochs_kept",
        hue="task",
        order=AGE_GROUPS,
        palette={"SuS": "#4E79A7", "CCD": "#E15759"},
        fliersize=1.2,
        linewidth=0.7,
        ax=axes[2],
    )
    axes[2].set_title("C Retained epochs by age")
    axes[2].set_xlabel("Age group")
    axes[2].set_xticklabels([AGE_LABELS[g] for g in AGE_GROUPS])
    axes[2].set_ylabel("Epochs")
    axes[2].legend(frameon=False, title="")
    axes[0].legend(frameon=False)
    axes[1].legend(frameon=False)
    sns.despine(fig)
    fig.tight_layout()
    out = FIGURES / "figS1"
    save_figure(fig, out)
    return out.with_suffix(".png")


def weighted_full_channel_component_means() -> tuple[pd.DataFrame, pd.DataFrame]:
    all_wf = pd.read_csv(ROOT / "results/tables/full_channel_erp_topomaps/full_channel_erp_all_channel_waveforms.csv")
    sample = pd.read_csv(ROOT / "results/tables/full_channel_erp_topomaps/full_channel_erp_sample_summary.csv")
    weights = sample.set_index(["component", "age_group"])["n_subjects"].to_dict()
    records = []
    for (component, age_group), df in all_wf.groupby(["component", "age_group"]):
        rec = df.copy()
        rec["weight"] = weights.get((component, age_group), 1)
        records.append(rec)
    all_wf = pd.concat(records, ignore_index=True)
    return all_wf, sample


def plot_s2_full_channel_task_waveforms() -> Path:
    all_wf, sample = weighted_full_channel_component_means()
    pivot = all_wf.pivot_table(
        index=["age_group", "channel", "time_s", "weight"], columns="component", values="amplitude_uv"
    ).reset_index()
    pivot["SuS"] = pivot["general"] + pivot["specific_sus_minus_ccd"] / np.sqrt(2)
    pivot["CCD"] = pivot["general"] - pivot["specific_sus_minus_ccd"] / np.sqrt(2)
    rows = []
    for (channel, time_s), df in pivot.groupby(["channel", "time_s"]):
        total_weight = df["weight"].sum()
        rows.append(
            {
                "channel": channel,
                "time_s": time_s,
                "SuS": np.average(df["SuS"], weights=df["weight"]),
                "CCD": np.average(df["CCD"], weights=df["weight"]),
            }
        )
    grand = pd.DataFrame(rows)
    grand.to_csv(TABLES / "figS2_task_level_full_channel_waveforms_source.csv", index=False)
    windows = [
        ("P1", 0.08, 0.14),
        ("P2", 0.14, 0.22),
        ("N450", 0.30, 0.60),
    ]
    fig = plt.figure(figsize=(7.7, 5.7))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.0, 1.18], hspace=0.42, wspace=0.28)
    ax_sus = fig.add_subplot(gs[0, 0])
    ax_ccd = fig.add_subplot(gs[0, 1], sharex=ax_sus, sharey=ax_sus)
    ax_overlay = fig.add_subplot(gs[1, 0], sharex=ax_sus)
    ax_diff = fig.add_subplot(gs[1, 1], sharex=ax_sus)
    times = np.sort(grand["time_s"].unique())
    channels = sorted(grand["channel"].unique(), key=lambda x: int(x[1:]) if x.startswith("E") and x[1:].isdigit() else 999)
    view = grand[grand["time_s"].between(-0.20, 0.80)]
    y_lim = max(4.0, min(12.0, float(np.nanpercentile(np.abs(view[["SuS", "CCD"]].to_numpy()), 99.5)) * 1.08))

    def mark_windows(ax, y_text: float | None = None, show_labels: bool = False) -> None:
        for idx, (label, start, end) in enumerate(windows):
            ax.axvspan(start, end, color="#D9A441", alpha=0.11 if idx < 2 else 0.15, linewidth=0)
            if show_labels and y_text is not None:
                ax.text((start + end) / 2, y_text, label, ha="center", va="top", fontsize=5.8, color="#3B3428")

    for task, ax, color, title in [
        ("SuS", ax_sus, "#4E79A7", "A SuS full-channel butterfly"),
        ("CCD", ax_ccd, "#E15759", "B CCD full-channel butterfly"),
    ]:
        mark_windows(ax, y_text=y_lim * 0.95, show_labels=True)
        for ch in channels:
            sub = grand[grand["channel"].eq(ch)].sort_values("time_s")
            ax.plot(sub["time_s"], sub[task], color=color, alpha=0.10, linewidth=0.42)
        mean = grand.groupby("time_s")[task].mean().reindex(times)
        ax.plot(times, mean, color=color, linewidth=1.55, label=f"{task} channel mean")
        ax.axvline(0, color="#666666", linewidth=0.6)
        ax.axhline(0, color="#BBBBBB", linewidth=0.6)
        ax.set_xlim(-0.20, 0.80)
        ax.set_ylim(-y_lim, y_lim)
        ax.set_title(title)
        ax.set_xlabel("Time (s)")
        ax.set_ylabel("Amplitude (uV)")
        ax.text(0.02, 0.06, "full-channel", transform=ax.transAxes, fontsize=5.8, color="#555555")

    for task, color in [("SuS", "#4E79A7"), ("CCD", "#E15759")]:
        mat = grand.pivot(index="channel", columns="time_s", values=task).reindex(channels)
        mat = mat.loc[:, times].to_numpy(dtype=float)
        gfp = np.sqrt(np.nanmean(mat**2, axis=0))
        ax_overlay.plot(times, gfp, color=color, linewidth=1.55, label=task)
    gfp_top = max(0.5, ax_overlay.get_ylim()[1])
    mark_windows(ax_overlay, y_text=gfp_top * 0.95, show_labels=True)
    ax_overlay.axvline(0, color="#666666", linewidth=0.6)
    ax_overlay.axhline(0, color="#BBBBBB", linewidth=0.6)
    ax_overlay.set_xlim(-0.20, 0.80)
    ax_overlay.set_ylim(0, gfp_top)
    ax_overlay.set_title("C Task overlay, global field power")
    ax_overlay.set_xlabel("Time (s)")
    ax_overlay.set_ylabel("GFP (uV)")
    ax_overlay.legend(frameon=False, loc="upper left", fontsize=6)

    diff = grand.pivot(index="channel", columns="time_s", values="SuS") - grand.pivot(
        index="channel", columns="time_s", values="CCD"
    )
    diff = diff.reindex(channels)
    vmax = np.nanpercentile(np.abs(diff.to_numpy()), 98)
    im = ax_diff.imshow(
        diff.to_numpy(),
        aspect="auto",
        origin="lower",
        extent=[times.min(), times.max(), 0, len(channels)],
        cmap="RdBu_r",
        vmin=-vmax,
        vmax=vmax,
    )
    ax_diff.set_title("D SuS - CCD by channel")
    ax_diff.set_xlabel("Time (s)")
    ax_diff.set_ylabel("Channel index")
    for label, start, end in windows:
        ax_diff.axvspan(start, end, color="black", alpha=0.08, linewidth=0)
        ax_diff.text((start + end) / 2, len(channels) - 2.0, label, ha="center", va="top", fontsize=6.2, color="#2F2F2F")
    fig.colorbar(im, ax=ax_diff, label="SuS - CCD (uV)", fraction=0.025, pad=0.02)
    sns.despine(fig)
    out = FIGURES / "figS2"
    save_figure(fig, out)
    return out.with_suffix(".png")


def mne_info_for_full_channel():
    import mne

    qc = pd.read_csv(ROOT / "results/tables/full_sus_ccd_epoch_qc_analysisch_reject250.csv")
    raw_path = Path(qc.loc[qc["status"].eq("ok"), "raw_path"].iloc[0])
    raw = mne.io.read_raw_eeglab(raw_path, preload=False, verbose="ERROR")
    raw.pick("eeg")
    return raw.info.copy(), raw.ch_names


def plot_s3_component_validation() -> Path:
    display_labels = {"P1": "P1", "N1": "P2", "P3": "N450"}
    selection = pd.read_csv(ROOT / "results/tables/selected_electrodes/selected_erp_component_electrodes.csv")
    selection = selection[selection["erp_component"].isin(["P1", "N1", "P3"])].copy()
    selection["channels_list"] = selection["channels"].map(lambda x: [ch.strip() for ch in str(x).split(",") if ch.strip()])
    selection = selection.set_index("erp_component")
    windows = {
        erp: (
            float(selection.loc[erp, "tmin_s"]),
            float(selection.loc[erp, "tmax_s"]),
            str(selection.loc[erp, "roi"]),
            selection.loc[erp, "channels_list"],
        )
        for erp in ["P1", "N1", "P3"]
    }
    values = pd.read_csv(ROOT / "results/tables/extended_full_channel_erp_components/extended_full_channel_erp_window_values.csv")
    waves = pd.read_csv(ROOT / "results/tables/selected_electrodes/selected_electrode_erp_group_waveforms.csv")
    sample = pd.read_csv(ROOT / "results/tables/extended_full_channel_erp_components/extended_full_channel_erp_sample_summary.csv")
    weights = sample[sample["component"].eq("general")].set_index("age_group")["n_subjects"].to_dict()
    info, ch_names = mne_info_for_full_channel()
    mean_maps = {}
    for erp, (_, _, _, channels) in windows.items():
        top = values[(values["component"].eq("general")) & (values["erp_component"].eq(erp))].copy()
        top["weight"] = top["age_group"].map(weights)
        mean_maps[erp] = top.groupby("channel").apply(
            lambda d: np.average(d["amplitude_uv"], weights=d["weight"])
        ).reindex(ch_names)
    all_map_vals = np.concatenate([v.to_numpy(dtype=float) for v in mean_maps.values()])
    vmax = max(1.0, float(np.nanpercentile(np.abs(all_map_vals[np.isfinite(all_map_vals)]), 98)))
    fig, axes = plt.subplots(
        2,
        3,
        figsize=(7.8, 4.2),
        gridspec_kw={"height_ratios": [1.04, 1.0], "hspace": 0.34, "wspace": 0.36},
    )
    im = None
    for col, (erp, (tmin, tmax, roi, channels)) in enumerate(windows.items()):
        label = display_labels.get(erp, erp)
        im = plot_topomap(mean_maps[erp].to_numpy(dtype=float), info, axes[0, col], vmax=vmax)
        highlight_roi(axes[0, col], info, ch_names, channels)
        axes[0, col].set_title(f"{label}\n{int(tmin*1000)}-{int(tmax*1000)} ms\n{roi.replace('_', ' ')}", fontsize=7)
        ax = axes[1, col]
        y_values = []
        for component, color, label, style, alpha in [
            ("general", "#333333", "General", "-", 1.0),
            ("specific_sus_minus_ccd", "#7462D8", "Specific", "--", 0.95),
        ]:
            wave = waves[
                (waves["component"].eq(component))
                & (waves["erp_component"].eq(erp))
                & (waves["roi"].eq(roi))
            ].copy()
            wave["weight"] = wave["age_group"].map(weights)
            agg = wave.groupby("time_s").apply(
                lambda d: np.average(d["amplitude_uv"], weights=d["weight"])
            ).reset_index(name="amplitude_uv")
            y_values.extend(agg["amplitude_uv"].to_numpy(dtype=float).tolist())
            ax.plot(agg["time_s"], agg["amplitude_uv"], color=color, linewidth=1.15, linestyle=style, alpha=alpha, label=label)
        ax = axes[1, col]
        ax.axvspan(tmin, tmax, color="#D9A441", alpha=0.18, linewidth=0)
        ax.axvline(0, color="#666666", linewidth=0.6)
        ax.axhline(0, color="#BBBBBB", linewidth=0.6)
        y = np.asarray(y_values, dtype=float)
        y = y[np.isfinite(y)]
        if len(y):
            lo, hi = np.nanpercentile(y, [1, 99])
            pad = max(0.25, (hi - lo) * 0.12)
            ax.set_ylim(lo - pad, hi + pad)
        ax.set_xlim(-0.20, 0.80)
        ax.set_title(f"Selected-ROI waveform", fontsize=7)
        ax.set_xlabel("Time (s)")
        if col == 0:
            ax.set_ylabel("Amplitude (uV)")
        if col == 2:
            ax.legend(frameon=False, fontsize=5.8, loc="upper left")
    cbar_ax = fig.add_axes([0.925, 0.58, 0.012, 0.24])
    fig.colorbar(im, cax=cbar_ax, label="Amplitude (uV)")
    sns.despine(fig)
    fig.subplots_adjust(left=0.075, right=0.90, top=0.88, bottom=0.12)
    out = FIGURES / "figS3"
    save_figure(fig, out)
    return out.with_suffix(".png")


def fit_quad_line(df: pd.DataFrame, y: str, ages: np.ndarray) -> tuple[np.ndarray, np.ndarray, np.ndarray]:
    model_df = df.dropna(subset=[y, "age", "sex"]).copy()
    model_df["age2"] = model_df["age"] ** 2
    model = smf.ols(f"{y} ~ age + age2 + C(sex)", data=model_df).fit()
    sex = model_df["sex"].mode().iloc[0]
    pred_df = pd.DataFrame({"age": ages, "age2": ages**2, "sex": sex})
    pred = model.get_prediction(pred_df).summary_frame(alpha=0.05)
    return pred["mean"].to_numpy(), pred["mean_ci_lower"].to_numpy(), pred["mean_ci_upper"].to_numpy()


def plot_s4_p1() -> Path:
    comp = load_components()
    df = comp[(comp["feature"].eq("p1")) & (comp["age_group"].isin(AGE_GROUPS))].copy()
    ages = np.linspace(5, 21, 160)
    fig, axes = plt.subplots(2, 3, figsize=(7.4, 4.8))
    for col, component in enumerate(COMPONENTS):
        ax = axes[0, col]
        for group in AGE_GROUPS:
            sub = df[df["age_group"].eq(group)]
            ax.scatter(sub["age"], sub[component], s=5, alpha=0.18, color=AGE_COLORS[group], rasterized=True)
        mean, lo, hi = fit_quad_line(df, component, ages)
        ax.plot(ages, mean, color=COMP_COLORS[component], linewidth=1.6)
        ax.fill_between(ages, lo, hi, color=COMP_COLORS[component], alpha=0.14, linewidth=0)
        ax.axhline(0, color="#BBBBBB", linewidth=0.6)
        ax.set_title(f"A{col+1} P1 {COMPONENT_LABELS[component]} age fit")
        ax.set_xlabel("Age (years)")
        ax.set_ylabel("Component score")
        ax = axes[1, col]
        sns.boxplot(data=df, x="age_group", y=component, order=AGE_GROUPS, palette=AGE_COLORS, fliersize=0.8, linewidth=0.7, ax=ax)
        ax.set_xticklabels([AGE_LABELS[g] for g in AGE_GROUPS])
        ax.set_title(f"B{col+1} P1 {COMPONENT_LABELS[component]} by age group")
        ax.set_xlabel("Age group")
        ax.set_ylabel("Component score")
    ax = axes[0, 2]
    anova = pd.read_csv(ROOT / "results/tables/full_sus_ccd_age_group_anova.csv")
    p1 = anova[(anova["feature"].eq("p1")) & (anova["term"].eq("C(age_group)"))].copy()
    vals = [float(p1.loc[p1["component"].eq(c), "r_squared"].iloc[0]) for c in COMPONENTS]
    ax.bar([0, 1], vals, color=[COMP_COLORS[c] for c in COMPONENTS], width=0.65)
    ax.set_xticks([0, 1])
    ax.set_xticklabels(["General", "Specific"])
    ax.set_ylabel("Age-model R2")
    ax.set_title("C P1 age effect size")
    axes[1, 2].axis("off")
    posthoc = pd.read_csv(ROOT / "results/tables/full_sus_ccd_followup/full_sus_ccd_age_group_posthoc.csv")
    p1post = posthoc[posthoc["feature"].eq("p1")][["component", "contrast", "t_value", "p_value", "q_value", "cohen_d_raw"]]
    text = "P1 post-hoc summary\n" + "\n".join(
        f"{r.component.replace('component_', '')}, {r.contrast}: q={fmt_p(r.q_value)}, d={r.cohen_d_raw:.2f}"
        for _, r in p1post.head(6).iterrows()
    )
    axes[1, 2].text(0, 0.95, text, va="top", ha="left", fontsize=6.5)
    sns.despine(fig)
    fig.tight_layout()
    out = FIGURES / "figS4"
    save_figure(fig, out)
    return out.with_suffix(".png")


def plot_s5_polynomial_fits() -> Path:
    comp = load_components()
    fig, axes = plt.subplots(3, 2, figsize=(7.4, 7.2), sharex=True)
    ages = np.linspace(5, 21, 160)
    source_rows = []
    for r, feature in enumerate(FEATURES):
        fdf = comp[(comp["feature"].eq(feature)) & (comp["age_group"].isin(AGE_GROUPS))].copy()
        for c, component in enumerate(COMPONENTS):
            ax = axes[r, c]
            for group in AGE_GROUPS:
                sub = fdf[fdf["age_group"].eq(group)]
                ax.scatter(sub["age"], sub[component], s=4, alpha=0.14, color=AGE_COLORS[group], rasterized=True)
            model_df = fdf.dropna(subset=[component, "age", "sex"]).copy()
            model_df["age2"] = model_df["age"] ** 2
            model = smf.ols(f"{component} ~ age + age2 + C(sex)", data=model_df).fit()
            pred_df = pd.DataFrame({"age": ages, "age2": ages**2, "sex": model_df["sex"].mode().iloc[0]})
            pred = model.get_prediction(pred_df).summary_frame(alpha=0.05)
            ax.plot(ages, pred["mean"], color=COMP_COLORS[component], linewidth=1.4)
            ax.fill_between(ages, pred["mean_ci_lower"], pred["mean_ci_upper"], color=COMP_COLORS[component], alpha=0.14, linewidth=0)
            ax.axhline(0, color="#BBBBBB", linewidth=0.6)
            ax.set_title(f"{FEATURE_LABELS[feature]} {COMPONENT_LABELS[component]}")
            ax.text(
                0.02,
                0.96,
                f"b_age={model.params.get('age', np.nan):.3f}\nb_age2={model.params.get('age2', np.nan):.3f}\nR2={model.rsquared:.3f}",
                transform=ax.transAxes,
                va="top",
                ha="left",
                fontsize=6,
            )
            ax.set_ylabel("Component score")
            source_rows.append(
                {
                    "feature": feature,
                    "component": component,
                    "n": int(model.nobs),
                    "age_beta": model.params.get("age", np.nan),
                    "age_p": model.pvalues.get("age", np.nan),
                    "age2_beta": model.params.get("age2", np.nan),
                    "age2_p": model.pvalues.get("age2", np.nan),
                    "r_squared": model.rsquared,
                }
            )
    axes[-1, 0].set_xlabel("Age (years)")
    axes[-1, 1].set_xlabel("Age (years)")
    pd.DataFrame(source_rows).to_csv(TABLES / "figS5_polynomial_fit_source.csv", index=False)
    sns.despine(fig)
    fig.tight_layout()
    out = FIGURES / "figS5"
    save_figure(fig, out)
    return out.with_suffix(".png")


def plot_s6_lme_diagnostics(diag: pd.DataFrame) -> Path:
    if diag.empty:
        raise ValueError("No mixed-model diagnostics available")
    fig, axes = plt.subplots(3, 3, figsize=(7.4, 7.0))
    for r, feature in enumerate(FEATURES):
        sub = diag[diag["feature"].eq(feature)].copy()
        sm.qqplot(sub["residual"], line="45", ax=axes[r, 0], markerfacecolor="#4E79A7", markeredgecolor="none", alpha=0.45)
        axes[r, 0].set_title(f"{FEATURE_LABELS[feature]} residual QQ")
        axes[r, 1].scatter(sub["fitted"], sub["residual"], s=5, alpha=0.22, color="#4E79A7", rasterized=True)
        axes[r, 1].axhline(0, color="#666666", linewidth=0.6)
        axes[r, 1].set_title("Residuals vs fitted")
        axes[r, 1].set_xlabel("Fitted")
        axes[r, 1].set_ylabel("Residual")
        re = sub.drop_duplicates("subject")
        axes[r, 2].hist(re["random_intercept"].dropna(), bins=28, alpha=0.72, color="#4E79A7", label="intercept")
        axes[r, 2].hist(re["random_slope"].dropna(), bins=28, alpha=0.45, color="#E15759", label="slope")
        axes[r, 2].set_title("Random effects")
        axes[r, 2].legend(frameon=False)
    sns.despine(fig)
    fig.tight_layout()
    out = FIGURES / "figS6"
    save_figure(fig, out)
    return out.with_suffix(".png")


def plot_s7_psych_distributions() -> Path:
    comp = load_components()
    subjects = comp[comp["feature"].eq("p1")].drop_duplicates("subject").dropna(subset=PSYCH).copy()
    fig = plt.figure(figsize=(7.4, 5.8))
    gs = fig.add_gridspec(2, 4, height_ratios=[1, 1.25])
    for i, psych in enumerate(PSYCH):
        ax = fig.add_subplot(gs[0, i])
        ax.hist(subjects[psych], bins=28, color="#4E79A7", alpha=0.75)
        ax.set_title(PSYCH_LABELS[psych])
        ax.set_xlabel("Score")
        if i == 0:
            ax.set_ylabel("Subjects")
    corr_ax = fig.add_subplot(gs[1, :2])
    corr = subjects[PSYCH].corr()
    sns.heatmap(corr, annot=True, fmt=".2f", cmap="vlag", center=0, vmin=-1, vmax=1, square=True, cbar_kws={"label": "r"}, ax=corr_ax)
    corr_ax.set_title("Correlation matrix")
    scatter_ax = fig.add_subplot(gs[1, 2:])
    scatter_ax.scatter(subjects["p_factor"], subjects["attention"], s=8, alpha=0.28, label="Attention", color="#4E79A7")
    scatter_ax.scatter(subjects["p_factor"], subjects["internalizing"], s=8, alpha=0.28, label="Internalizing", color="#59A14F")
    scatter_ax.scatter(subjects["p_factor"], subjects["externalizing"], s=8, alpha=0.28, label="Externalizing", color="#E15759")
    scatter_ax.axhline(0, color="#BBBBBB", linewidth=0.6)
    scatter_ax.axvline(0, color="#BBBBBB", linewidth=0.6)
    scatter_ax.set_xlabel("p-factor")
    scatter_ax.set_ylabel("Specific dimension score")
    scatter_ax.set_title("Dimension scatter")
    scatter_ax.legend(frameon=False)
    subjects[["subject", *PSYCH]].to_csv(TABLES / "figS7_psychopathology_distribution_source.csv", index=False)
    sns.despine(fig)
    fig.tight_layout()
    out = FIGURES / "figS7"
    save_figure(fig, out)
    return out.with_suffix(".png")


def plot_s8_pfactor_tertiles() -> Path:
    comp = load_components()
    comp = comp.dropna(subset=["p_factor", "age", "sex"]).copy()
    comp["p_tertile"] = pd.qcut(comp["p_factor"], 3, labels=["Low", "Middle", "High"])
    ages = np.linspace(5, 21, 150)
    fig, axes = plt.subplots(3, 2, figsize=(7.4, 7.2), sharex=True)
    tertile_colors = {"Low": "#4E79A7", "Middle": "#8A8A8A", "High": "#E15759"}
    source_rows = []
    for r, feature in enumerate(FEATURES):
        fdf = comp[comp["feature"].eq(feature)].copy()
        for c, component in enumerate(COMPONENTS):
            ax = axes[r, c]
            for tertile in ["Low", "Middle", "High"]:
                tdf = fdf[fdf["p_tertile"].eq(tertile)].dropna(subset=[component]).copy()
                if len(tdf) < 20:
                    continue
                tdf["age2"] = tdf["age"] ** 2
                model = smf.ols(f"{component} ~ age + age2 + C(sex)", data=tdf).fit()
                pred_df = pd.DataFrame({"age": ages, "age2": ages**2, "sex": tdf["sex"].mode().iloc[0]})
                pred = model.get_prediction(pred_df).summary_frame(alpha=0.05)
                ax.plot(ages, pred["mean"], color=tertile_colors[tertile], linewidth=1.2, label=tertile if r == 0 and c == 0 else None)
                ax.fill_between(ages, pred["mean_ci_lower"], pred["mean_ci_upper"], color=tertile_colors[tertile], alpha=0.09, linewidth=0)
                for age, y, lo, hi in zip(ages, pred["mean"], pred["mean_ci_lower"], pred["mean_ci_upper"]):
                    if abs((age * 10) % 10) < 1e-6:
                        source_rows.append(
                            {
                                "feature": feature,
                                "component": component,
                                "p_tertile": tertile,
                                "age": age,
                                "predicted": y,
                                "ci_low": lo,
                                "ci_high": hi,
                                "n": int(model.nobs),
                            }
                        )
            ax.axhline(0, color="#BBBBBB", linewidth=0.6)
            ax.set_title(f"{FEATURE_LABELS[feature]} {COMPONENT_LABELS[component]}")
            ax.set_ylabel("Component score")
    axes[0, 0].legend(title="p-factor", frameon=False, loc="best")
    axes[-1, 0].set_xlabel("Age (years)")
    axes[-1, 1].set_xlabel("Age (years)")
    pd.DataFrame(source_rows).to_csv(TABLES / "figS8_pfactor_tertile_predictions_source.csv", index=False)
    sns.despine(fig)
    fig.tight_layout()
    out = FIGURES / "figS8"
    save_figure(fig, out)
    return out.with_suffix(".png")


def plot_s9_ccd_behavior() -> Path:
    beh = pd.read_csv(ROOT / "results/tables/clinical_rescue_behavior/ccd_behavior_by_subject.csv")
    meta = load_components()[["subject", "age", "age_group", "sex", "p_factor"]].drop_duplicates()
    df = beh.merge(meta, on="subject", how="left")
    df = df[df["age_group"].isin(AGE_GROUPS)].copy()
    df["log_rt"] = np.log(pd.to_numeric(df["ccd_median_rt_correct"], errors="coerce"))
    fig, axes = plt.subplots(1, 3, figsize=(7.4, 2.5))
    sns.boxplot(data=df, x="age_group", y="ccd_accuracy", order=AGE_GROUPS, palette=AGE_COLORS, fliersize=1, linewidth=0.7, ax=axes[0])
    axes[0].set_title("A CCD accuracy")
    axes[0].set_xticklabels([AGE_LABELS[g] for g in AGE_GROUPS])
    axes[0].set_xlabel("Age group")
    axes[0].set_ylabel("Accuracy")
    sns.boxplot(data=df, x="age_group", y="ccd_median_rt_correct", order=AGE_GROUPS, palette=AGE_COLORS, fliersize=1, linewidth=0.7, ax=axes[1])
    axes[1].set_title("B Correct-trial RT")
    axes[1].set_xticklabels([AGE_LABELS[g] for g in AGE_GROUPS])
    axes[1].set_xlabel("Age group")
    axes[1].set_ylabel("Median RT (s)")
    axes[2].scatter(df["p_factor"], df["ccd_accuracy"], s=8, alpha=0.35, color="#4E79A7", label="Accuracy")
    ax2 = axes[2].twinx()
    ax2.scatter(df["p_factor"], df["ccd_median_rt_correct"], s=8, alpha=0.22, color="#E15759", label="RT")
    axes[2].set_title("C Behaviour vs p-factor")
    axes[2].set_xlabel("p-factor")
    axes[2].set_ylabel("Accuracy", color="#4E79A7")
    ax2.set_ylabel("Median RT (s)", color="#E15759")
    df.to_csv(TABLES / "figS9_CCD_behavior_source.csv", index=False)
    sns.despine(fig, right=False)
    fig.tight_layout()
    out = FIGURES / "figS9"
    save_figure(fig, out)
    return out.with_suffix(".png")


def supplementary_texts() -> dict[str, str]:
    s1 = """Supplementary Text S1. Detailed description of the HBN task paradigms.

The surround-suppression (SuS) task was a passive visual paradigm. Each trial began with a fixation spot for 500 ms, followed by simultaneous presentation of the foreground and surround stimuli for 2,400 ms and a 500-ms inter-trial interval. The foreground consisted of four vertical sinusoidal luminance-modulated grating disks with a radius of 2 degrees and spatial frequency of 1 cycle per degree. The disks were placed at 5 degrees eccentricity, at polar angles 20 degrees above and 45 degrees below the horizontal meridian. The foreground flickered on and off at 25 Hz, with upper and lower disks in opposite temporal phase. Across trials, foreground contrast varied across four levels and the surround was absent or present at full contrast, with surround orientation either parallel or orthogonal to the foreground. Two 64-trial blocks were collected. In the current analysis, SuS epochs were locked to the BIDS/HED value stim_ON, corresponding to the onset of the foreground/surround stimulation period.

The contrast-change-detection (CCD) task was an active visual decision paradigm. Participants viewed an annular stimulus with an inner radius of 1 degree and outer radius of 6 degrees. The annulus contained two overlaid gratings tilted 45 degrees leftward and rightward from vertical, phase-reversing at 20 and 25 Hz. At baseline, both gratings had 50% contrast. On target trials, one grating linearly increased from 50% to 100% contrast while the other decreased from 50% to 0% over 1,600 ms, followed by an 800-ms return to baseline. Participants used left- and right-hand responses to identify which tilted grating increased. Each block contained 12 left-tilted and 12 right-tilted targets in random order, with inter-target intervals of 2.8, 4.4, or 6.0 s. Feedback was presented for 400 ms after target end.

Event extraction used the task-specific events.tsv files and BIDS/HED sidecar definitions. SuS epochs were locked to stim_ON events. CCD epochs were locked to contrastTrial_start events, which mark the start of a single contrast-change trial and precede the left_target/right_target contrast-change events. Thus, the two tasks were aligned to visual/trial-onset events rather than to motor responses or feedback. At the physical level, SuS visual onset reflects onset of flickering peripheral disks embedded in a surround context, whereas CCD trial onset occurs within a continuously viewed annular grating display. This is why the decomposed shared signal is interpreted as a task-general visual-onset/trial-onset component across two visual tasks, rather than a response to physically identical stimulation."""
    s2 = """Supplementary Text S2. Rationale for preprocessing choices.

The primary ERP analysis used a 1-40 Hz band-pass filter, 60 Hz notch filter, average reference, -200 to 800 ms epochs, baseline correction to the pre-stimulus interval, and a 250 microvolt peak-to-peak rejection threshold. The 1-Hz high-pass cutoff reduced slow drift in long developmental recordings while preserving the P1, P2 and N450 time range. The 40-Hz low-pass cutoff was appropriate because the confirmatory outcomes were time-domain ERP components rather than high-frequency oscillatory responses, and it reduced high-frequency muscle and line-related residuals. A 60-Hz notch filter was used because the recordings were collected in a 60-Hz power-line environment.

The 250 microvolt threshold was chosen as a pragmatic compromise for a large child and adolescent EEG dataset. More stringent thresholds can improve single-epoch cleanliness but risk disproportionate data loss in younger children. The threshold retained 214,393 of 234,603 candidate epochs (91.4%) while excluding large-amplitude artifacts. Rejection was applied to the predefined analysis channels used for the ERP features, keeping the main component estimates protected from large excursions in the channels entering the analysis.

ICA sensitivity outputs are provided as robustness checks when available. The open-source package keeps the primary amplitude-based QC and the ICA sensitivity workflow separate so that users can reproduce either analysis path transparently."""
    s3 = """Supplementary Text S3. LME technical details and alternative model comparison.

The main task-general and task-specific variables were computed from task-z-scored SuS and CCD amplitudes as the within-participant mean and SuS-minus-CCD deviation. As an audit of this decomposition, we additionally fitted mixed-effects models separately for P1, P2 and N450 using the participant-by-task table. The confirmatory mixed model was amplitude_z ~ task + (1 + task | subject), with a simpler random-intercept-only model amplitude_z ~ task + (1 | subject) used as a comparison. Models were fitted by ML for AIC/BIC and likelihood-ratio comparison and by REML for final variance-component inspection.

The mixed-model audit is reported in Supplementary Table S2. The random-intercept plus random-slope models converged for the three ERP features, although singular-covariance warnings occurred in some fits, as expected when each participant contributes two task observations. The LME audit therefore serves as a technical robustness check rather than a replacement for the transparent two-task mean/deviation decomposition used in the main analysis. Model diagnostics based on residuals, fitted values and random-effect distributions are shown in Supplementary Figure S6."""
    return {"S1": s1, "S2": s2, "S3": s3}


TABLE_LEGENDS = {
    "S1": "Robustness analysis summary across the full sample, high-quality EEG subsample, complete-phenotype subsample, and their intersection.",
    "S2": "Mixed-effects variance-decomposition audit, including convergence, fixed task effects, random-effect variances, residual variance, AIC/BIC, likelihood-ratio tests and bootstrap ICC intervals.",
    "S3": "Complete age-group post-hoc pairwise comparisons for all ERP features and decomposition components.",
    "S4": "Split-half reliability for task-level ERPs and decomposed component scores. The resampling columns summarize 1000 subject-level bootstrap resamples of the available odd/even split.",
    "S5": "Complete psychopathology analyses, including partial associations for four CBCL-derived dimensions and age-by-dimension interaction tests.",
    "S6": "Trial-to-trial variability associations with p-factor, attention, internalizing and externalizing dimensions.",
    "S7": "CCD behavioural p-factor models and behavioural mediation screen.",
    "S8": "Sample demographic and data-quality characteristics by age group.",
}


FIGURE_LEGENDS = {
    "S1": "Preprocessing quality control. (A) Retained epoch counts for SuS and CCD. (B) Epoch rejection distributions. (C) Age-group-stratified retained epoch counts.",
    "S2": "Full-channel grand-average ERP waveforms. (A) SuS butterfly plot. (B) CCD butterfly plot. (C) SuS and CCD global field power overlay. (D) Channel-by-time SuS minus CCD difference heat map. P1, P2 and N450 measurement windows are shaded.",
    "S3": "ERP component extraction validation for the selected-electrode analyses. Top row shows full-channel task-general topographies for P1, P2 and N450 with the selected ROI electrodes marked. Bottom row shows the corresponding selected-ROI general and specific waveforms with measurement windows shaded.",
    "S4": "Complete P1 developmental results, including continuous-age fits, age-group distributions and age-effect size comparison.",
    "S5": "Continuous-age polynomial developmental fits for P1, P2 and N450 general and specific components.",
    "S6": "Mixed-effects model diagnostics: residual QQ plots, residuals versus fitted values and random-effect distributions.",
    "S7": "Psychopathology dimension distributions and correlation structure in the analysis sample.",
    "S8": "Complete p-factor tertile developmental curves for all six ERP component scores.",
    "S9": "CCD behavioural performance across development and associations with p-factor.",
}


def build_final_docx(paths: dict[str, Path]) -> Path:
    doc = Document()
    doc.add_heading("Supplementary Materials", level=0)
    for text in supplementary_texts().values():
        title, *body = text.split("\n\n")
        doc.add_heading(title, level=1)
        for para in body:
            p = doc.add_paragraph()
            run = p.add_run(para)
            run.font.size = Pt(10)
    doc.add_heading("Supplementary Tables", level=1)
    for key in [f"S{i}" for i in range(1, 9)]:
        doc.add_heading(f"Supplementary Table {key}.", level=2)
        p = doc.add_paragraph(TABLE_LEGENDS[key])
        p.runs[0].font.size = Pt(10)
        p2 = doc.add_paragraph(f"Source file: {paths[f'Table {key}']}")
        p2.runs[0].font.size = Pt(9)
    doc.add_heading("Supplementary Figures", level=1)
    for key in [f"S{i}" for i in range(1, 10)]:
        doc.add_heading(f"Supplementary Figure {key}.", level=2)
        fig_path = paths[f"Figure {key}"]
        if fig_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(fig_path), width=Inches(6.4))
        cap = doc.add_paragraph(FIGURE_LEGENDS[key])
        cap.runs[0].font.size = Pt(9)
    out = OUT / "Supplementary_Materials_Final.docx"
    doc.save(out)
    return out


def main() -> None:
    ensure_dirs()
    configure_style()
    paths: dict[str, Path] = {}
    paths["Table S1"] = write_table_s1()
    paths["Table S2"], diag = write_table_s2_and_diagnostics()
    paths["Table S3"] = write_table_s3()
    paths["Table S4"] = write_table_s4()
    s5a, s5b = write_table_s5()
    paths["Table S5"] = s5a
    paths["Table S5B"] = s5b
    paths["Table S6"] = write_table_s6()
    s7a, s7b = write_table_s7()
    paths["Table S7"] = s7a
    paths["Table S7B"] = s7b
    paths["Table S8"] = write_table_s8()

    paths["Figure S1"] = plot_s1_qc()
    paths["Figure S2"] = plot_s2_full_channel_task_waveforms()
    paths["Figure S3"] = plot_s3_component_validation()
    paths["Figure S4"] = plot_s4_p1()
    paths["Figure S5"] = plot_s5_polynomial_fits()
    paths["Figure S6"] = plot_s6_lme_diagnostics(diag)
    paths["Figure S7"] = plot_s7_psych_distributions()
    paths["Figure S8"] = plot_s8_pfactor_tertiles()
    paths["Figure S9"] = plot_s9_ccd_behavior()

    for key, text in supplementary_texts().items():
        (TEXT / f"Supplementary_Text_{key}.md").write_text(text + "\n", encoding="utf-8")

    inventory_rows = [{"item": key, "path": str(value)} for key, value in paths.items()]
    inventory = pd.DataFrame(inventory_rows)
    inventory.to_csv(OUT / "supplementary_final_inventory.csv", index=False)
    final_docx = build_final_docx(paths)
    print(final_docx)
    print(OUT / "supplementary_final_inventory.csv")
    print(f"tables={len(list(TABLES.glob('*.csv')))} figures={len(list(FIGURES.glob('*.png')))}")


if __name__ == "__main__":
    main()
