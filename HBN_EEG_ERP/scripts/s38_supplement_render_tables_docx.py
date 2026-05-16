from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUPP_ROOT = ROOT / "results" / "supplementary_materials_final"
TABLE_DIR = SUPP_ROOT / "tables"
FINAL_DOCX = SUPP_ROOT / "Supplementary_Materials_Final.docx"
TABLE_DOCX = SUPP_ROOT / "Supplementary_Tables_Rendered.docx"
COMBINED_DOCX = SUPP_ROOT / "Supplementary_Materials_Final_with_tables.docx"
INVENTORY = SUPP_ROOT / "supplementary_final_inventory.csv"


TABLE_FILES = {
    "S1": "Supplementary_Table_S1_robustness_all_core_effects.csv",
    "S2": "Supplementary_Table_S2_LME_variance_decomposition_audit.csv",
    "S3": "Supplementary_Table_S3_all_posthoc_pairwise_comparisons.csv",
    "S4": "Supplementary_Table_S4_split_half_reliability_complete.csv",
    "S5A": "Supplementary_Table_S5A_psychopathology_partial_associations_complete.csv",
    "S5B": "Supplementary_Table_S5B_psychopathology_age_interactions_complete.csv",
    "S6": "Supplementary_Table_S6_IIV_four_dimension_associations_complete.csv",
    "S7A": "Supplementary_Table_S7A_CCD_behavior_pfactor_models.csv",
    "S7B": "Supplementary_Table_S7B_CCD_behavior_mediation_models.csv",
    "S8": "Supplementary_Table_S8_demographics_and_epoch_counts_complete.csv",
}


CAPTIONS = {
    "S1": "Supplementary Table S1. Robustness analysis summary across analytic samples.",
    "S2A": "Supplementary Table S2A. LME model fit and convergence audit.",
    "S2B": "Supplementary Table S2B. LME fixed-effect estimates and variance components.",
    "S2C": "Supplementary Table S2C. ICC estimates and random-slope model comparison.",
    "S3": "Supplementary Table S3. Complete post hoc pairwise comparisons across age groups.",
    "S4": "Supplementary Table S4. Complete split-half reliability estimates.",
    "S5A": "Supplementary Table S5A. Psychopathology dimension associations with ERP component scores.",
    "S5B": "Supplementary Table S5B. Psychopathology-by-age interaction models.",
    "S6": "Supplementary Table S6. Trial-to-trial variability associations with psychopathology dimensions.",
    "S7A": "Supplementary Table S7A. Associations between p-factor and CCD behavioral measures.",
    "S7B": "Supplementary Table S7B. CCD behavioral mediation analysis results.",
    "S8": "Supplementary Table S8. Demographic and epoch-retention characteristics by age group.",
}


NOTES = {
    "S1": "Core developmental effects are reported for the full component sample and three sensitivity samples. q values are FDR-adjusted within this supplementary table.",
    "S2A": "The random-intercept and random-intercept-plus-task-slope models were fitted using ML and REML. AIC/BIC are only defined for ML fits in statsmodels.",
    "S2B": "Task was coded as SuS versus CCD. Variance components are reported on the z-scored component-amplitude scale used for model auditing.",
    "S2C": "ICC values index the task-general subject-level component. Confidence intervals were obtained by bootstrap resampling where available.",
    "S3": "Contrasts compare age groups for each component estimate. q values are FDR-adjusted across the complete set of post hoc tests.",
    "S4": "Reliability estimates use the available odd/even split-half estimates; the resampling columns summarize subject-level bootstrap resamples of that split.",
    "S5A": "Models adjusted for age and sex when these covariates were available. No association survived FDR correction.",
    "S5B": "Interaction terms test whether psychopathology dimensions moderated linear or quadratic age-related ERP trajectories. No interaction survived FDR correction.",
    "S6": "IIV denotes intra-individual trial-to-trial ERP variability. No IIV association survived FDR correction.",
    "S7A": "Behavioral models summarize CCD accuracy and reaction-time associations with the p-factor.",
    "S7B": "Indirect effects are reported with bootstrap confidence intervals when available.",
    "S8": "Race/ethnicity was not available in the local HBN EEG participant files used for this analysis and is therefore noted as unavailable.",
}


DISPLAY_NAMES = {
    "sample": "Sample",
    "feature": "Feature",
    "component": "Component",
    "n": "n",
    "n_observations": "Observations",
    "n_subjects": "Subjects",
    "df": "df",
    "df_diff": "df diff.",
    "f_value": "F",
    "p_value": "p",
    "q_value": "q",
    "r_squared": "R2",
    "model_r_squared": "Model R2",
    "delta_r2": "Delta R2",
    "model": "Model",
    "estimation": "Estimation",
    "converged": "Converged",
    "warning_or_error": "Warning/error",
    "optimizer": "Optimizer",
    "log_likelihood": "Log likelihood",
    "aic": "AIC",
    "bic": "BIC",
    "fixed_intercept_beta": "Intercept beta",
    "fixed_intercept_se": "Intercept SE",
    "fixed_task_beta": "Task beta",
    "fixed_task_se": "Task SE",
    "subject_intercept_var": "Subject intercept var.",
    "subject_task_slope_var": "Subject x task slope var.",
    "residual_var": "Residual var.",
    "general_icc": "General ICC",
    "general_icc_ci_low": "ICC CI low",
    "general_icc_ci_high": "ICC CI high",
    "model_comparison": "Comparison",
    "lr_statistic": "LR statistic",
    "lr_p_value": "LR p",
    "contrast": "Contrast",
    "group_a": "Group A",
    "group_b": "Group B",
    "n_a": "n A",
    "n_b": "n B",
    "mean_a": "Mean A",
    "mean_b": "Mean B",
    "adjusted_difference": "Adjusted difference",
    "std_error": "SE",
    "t_value": "t",
    "cohens_d": "Cohen's d",
    "estimate_family": "Estimate family",
    "task": "Task",
    "raw_split_half_r": "Raw split-half r",
    "spearman_brown": "Spearman-Brown",
    "resampling_mean_spearman_brown": "Resampling mean",
    "resampling_sd_spearman_brown": "Resampling SD",
    "resampling_note": "Resampling note",
    "dimension": "Dimension",
    "estimate": "Estimate",
    "std_error": "SE",
    "partial_r": "Partial r",
    "test": "Test",
    "age_x_dimension_beta": "Age x dimension beta",
    "age2_x_dimension_beta": "Age2 x dimension beta",
    "iiv_outcome": "IIV outcome",
    "mediator": "Mediator",
    "outcome_modeled": "Outcome modeled",
    "p_factor_estimate": "p-factor beta",
    "p_factor_se": "p-factor SE",
    "p_factor_t": "p-factor t",
    "p_factor_p": "p-factor p",
    "r2": "R2",
    "outcome": "Outcome",
    "a_pfactor_to_mediator": "a: p-factor to mediator",
    "a_p_value": "a p",
    "b_mediator_to_erp": "b: mediator to ERP",
    "b_p_value": "b p",
    "direct_pfactor_to_erp": "Direct effect",
    "direct_p_value": "Direct p",
    "total_pfactor_to_erp": "Total effect",
    "total_p_value": "Total p",
    "indirect_effect": "Indirect effect",
    "bootstrap_ci_low": "Bootstrap CI low",
    "bootstrap_ci_high": "Bootstrap CI high",
    "bootstrap_p": "Bootstrap p",
    "age_group": "Age group",
    "n_subjects": "Subjects",
    "n_female": "Female",
    "n_male": "Male",
    "race_ethnicity": "Race/ethnicity",
    "n_full_phenotype": "Full phenotype n",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, font_size: float, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def configure_landscape_section(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def setup_document(doc: Document) -> None:
    configure_landscape_section(doc.sections[-1])

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(11)


def title_case_label(name: str) -> str:
    return DISPLAY_NAMES.get(name, name.replace("_", " ").capitalize())


def format_value(value, column: str) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"

    col = column.lower()
    if "warning" in col or "note" in col or col in {"sample", "feature", "component", "model", "estimation", "optimizer", "contrast", "group_a", "group_b", "estimate_family", "task", "dimension", "test", "iiv_outcome", "mediator", "outcome", "outcome_modeled", "age_group", "race_ethnicity", "model_comparison"}:
        return str(value)

    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)

    if col.startswith("n") or col in {"df", "df_diff"}:
        return f"{int(round(number))}"
    if col.endswith("p") or "p_value" in col or col in {"p", "q", "q_value", "lr_p_value", "bootstrap_p", "p_factor_p"}:
        if number < 0.001:
            return "<0.001"
        return f"{number:.3f}"
    if "aic" in col or "bic" in col or "log_likelihood" in col:
        return f"{number:.1f}"
    if "f_value" in col or col == "f" or "t_value" in col or col.endswith("_t") or "lr_statistic" in col:
        return f"{number:.2f}"
    if "r_squared" in col or col in {"r2", "delta_r2"}:
        return f"{number:.3f}"
    return f"{number:.3f}"


def display_frame(df: pd.DataFrame, columns: Iterable[str] | None = None) -> pd.DataFrame:
    if columns is not None:
        cols = [c for c in columns if c in df.columns]
        df = df.loc[:, cols].copy()
    else:
        df = df.copy()
    for col in df.columns:
        df[col] = df[col].map(lambda value, c=col: format_value(value, c))
    df.columns = [title_case_label(c) for c in df.columns]
    return df


def add_caption(doc: Document, caption: str, note: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(caption)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    if note:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run("Note. " + note)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(8)
        r2.italic = True


def add_word_table(doc: Document, df: pd.DataFrame, font_size: float = 7.0, repeat_header: bool = True) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.rows[0]
    for idx, column in enumerate(df.columns):
        cell = header.cells[idx]
        set_cell_text(cell, str(column), font_size, bold=True)
        set_cell_shading(cell, "D9EAF7")

    if repeat_header:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), font_size)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for m in ["top", "left", "bottom", "right"]:
                node = tc_mar.find(qn(f"w:{m}"))
                if node is None:
                    node = OxmlElement(f"w:{m}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "36")
                node.set(qn("w:type"), "dxa")


def add_page_break(doc: Document) -> None:
    if len(doc.paragraphs) > 2:
        doc.add_page_break()


def load_table(key: str) -> pd.DataFrame:
    return pd.read_csv(TABLE_DIR / TABLE_FILES[key])


def add_s1(doc: Document) -> None:
    add_page_break(doc)
    add_caption(doc, CAPTIONS["S1"], NOTES["S1"])
    df = display_frame(load_table("S1"))
    add_word_table(doc, df, font_size=7.0)


def add_s2(doc: Document) -> None:
    df = load_table("S2")
    panels = [
        (
            "S2A",
            [
                "feature",
                "model",
                "estimation",
                "n_observations",
                "n_subjects",
                "converged",
                "optimizer",
                "log_likelihood",
                "aic",
                "bic",
                "warning_or_error",
            ],
            6.5,
        ),
        (
            "S2B",
            [
                "feature",
                "model",
                "estimation",
                "fixed_intercept_beta",
                "fixed_intercept_se",
                "fixed_task_beta",
                "fixed_task_se",
                "subject_intercept_var",
                "subject_task_slope_var",
                "residual_var",
            ],
            6.7,
        ),
        (
            "S2C",
            [
                "feature",
                "model",
                "estimation",
                "general_icc",
                "general_icc_ci_low",
                "general_icc_ci_high",
                "model_comparison",
                "lr_statistic",
                "lr_p_value",
            ],
            7.0,
        ),
    ]
    for key, columns, font_size in panels:
        add_page_break(doc)
        add_caption(doc, CAPTIONS[key], NOTES[key])
        add_word_table(doc, display_frame(df, columns), font_size=font_size)


def add_standard_table(doc: Document, key: str, font_size: float = 7.0, columns: Iterable[str] | None = None) -> None:
    add_page_break(doc)
    add_caption(doc, CAPTIONS[key], NOTES[key])
    add_word_table(doc, display_frame(load_table(key), columns), font_size=font_size)


def add_s8(doc: Document) -> None:
    add_page_break(doc)
    add_caption(doc, CAPTIONS["S8"], NOTES["S8"])
    df = load_table("S8")
    age_groups = list(df["age_group"])
    rows = []
    for col in df.columns:
        if col == "age_group":
            continue
        row = {"Measure": title_case_label(col)}
        for age_group in age_groups:
            value = df.loc[df["age_group"] == age_group, col].iloc[0]
            row[str(age_group)] = format_value(value, col)
        rows.append(row)
    out = pd.DataFrame(rows)
    add_word_table(doc, out, font_size=7.2)


def add_all_tables(doc: Document) -> None:
    add_s1(doc)
    add_s2(doc)
    add_standard_table(doc, "S3", font_size=6.4)
    add_standard_table(doc, "S4", font_size=6.5)
    add_standard_table(doc, "S5A", font_size=6.7)
    add_standard_table(doc, "S5B", font_size=6.2)
    add_standard_table(doc, "S6", font_size=6.7)
    add_standard_table(doc, "S7A", font_size=7.0)
    add_standard_table(doc, "S7B", font_size=5.9)
    add_s8(doc)


def append_document_body(dst: Document, src: Document) -> None:
    for element in src.element.body:
        if element.tag.endswith("sectPr"):
            continue
        dst.element.body.append(deepcopy(element))


def update_inventory() -> None:
    lines = []
    if INVENTORY.exists():
        lines = INVENTORY.read_text(encoding="utf-8").splitlines()
    additions = [
        "Rendered Supplementary Tables DOCX,results/supplementary_materials_final/Supplementary_Tables_Rendered.docx",
        "Supplementary Materials With Rendered Tables DOCX,results/supplementary_materials_final/Supplementary_Materials_Final_with_tables.docx",
    ]
    for line in additions:
        if line not in lines:
            lines.append(line)
    INVENTORY.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    table_doc = Document()
    setup_document(table_doc)
    title = table_doc.add_heading("Supplementary Tables", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = table_doc.add_paragraph(
        "Rendered table version for direct manuscript/supplement review. Numeric values are rounded for display; machine-readable source CSV files remain available in the supplementary table directory."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_all_tables(table_doc)
    table_doc.save(TABLE_DOCX)

    combined = Document(FINAL_DOCX)
    configure_landscape_section(combined.add_section(WD_SECTION.NEW_PAGE))
    combined.add_heading("Rendered Supplementary Tables", level=1)
    intro2 = combined.add_paragraph(
        "Rendered table version for direct manuscript/supplement review. Numeric values are rounded for display; machine-readable source CSV files remain available in the supplementary table directory."
    )
    intro2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_all_tables(combined)
    combined.save(COMBINED_DOCX)

    update_inventory()
    print(f"Saved rendered tables: {TABLE_DOCX}")
    print(f"Saved combined supplement: {COMBINED_DOCX}")


if __name__ == "__main__":
    main()
